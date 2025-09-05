"""
Modular conversation parser system for different chat UIs
Each company's chat export has different formatting that needs specific handling
"""

from abc import ABC, abstractmethod
from typing import List, Dict, Optional
import json
import re
import requests
from urllib.parse import urlparse
import time
try:
    from docx import Document
    from docx.shared import RGBColor
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False


class BaseConversationParser(ABC):
    """
    Abstract base class for conversation parsers
    Each company's parser should extend this
    """
    
    @abstractmethod
    def can_parse(self, content: str) -> bool:
        """
        Check if this parser can handle the given content
        Returns True if the format matches this parser's expected format
        """
        pass
    
    @abstractmethod
    def parse(self, content: str) -> List[Dict]:
        """
        Parse the conversation content into a standardized format
        Returns list of messages with 'role' and 'content' keys
        """
        pass
    
    @abstractmethod
    def get_parser_name(self) -> str:
        """Return the name of this parser for UI display"""
        pass
    
    @abstractmethod
    def get_company_name(self) -> str:
        """Return the company name this parser is for"""
        pass
    
    def clean_content(self, content: str) -> str:
        """
        Clean up content - remove extra whitespace, fix encoding issues, etc.
        Can be overridden by specific parsers
        """
        # Remove zero-width spaces and other invisible characters
        content = re.sub(r'[\u200b\u200c\u200d\ufeff]', '', content)
        # Normalize whitespace
        content = re.sub(r'\n{3,}', '\n\n', content)
        return content.strip()


class AnthropicChatParser(BaseConversationParser):
    """
    Parser for Anthropic Claude chat exports via Word documents
    
    Expected format:
    - Word document (.docx) with conversation pasted from Claude.ai
    - Highlighted text = User messages
    - Non-highlighted text = Assistant messages
    """
    
    def can_parse(self, content: str) -> bool:
        """Check if this looks like a Word document path or if we should handle it"""
        # This parser is selected explicitly when user chooses Anthropic
        # and uploads a .docx file, so we can be more liberal here
        return True  # Will be validated in parse() method
    
    def parse(self, content: str) -> List[Dict]:
        """
        Parse Anthropic conversation from Word document
        Content should be the file path to a .docx file
        """
        if not DOCX_AVAILABLE:
            return [{
                "role": "system",
                "content": "Error: python-docx library not available. Please install it with: pip install python-docx"
            }]
        
        try:
            return self._parse_docx_file(content)
        except Exception as e:
            return [{
                "role": "system", 
                "content": f"Error parsing Word document: {str(e)}"
            }]
    
    def _parse_docx_file(self, file_path: str) -> List[Dict]:
        """
        Parse Word document with different highlight colors for user vs LLM messages
        
        Logic:
        - Two different highlight colors distinguish user vs assistant messages
        - Group consecutive paragraphs with the same highlight color into single messages
        - Automatically detect which color represents user vs assistant based on pattern
        """
        try:
            doc = Document(file_path)
        except Exception as e:
            raise Exception(f"Could not open Word document: {str(e)}")
        
        # First pass: collect all highlight colors and their frequencies
        color_analysis = self._analyze_highlight_colors(doc)
        
        # Determine which colors represent user vs assistant
        color_mapping = self._determine_color_roles(color_analysis)
        
        messages = []
        current_message = {"role": None, "content": [], "color": None}
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue  # Skip empty paragraphs
            
            # Get the highlight color for this paragraph
            paragraph_color = self._get_paragraph_highlight_color(paragraph)
            
            # Determine role based on color
            role = color_mapping.get(paragraph_color, "assistant")  # Default to assistant
            
            # If color/role changes, save current message and start new one
            if (current_message["role"] is not None and 
                (current_message["role"] != role or current_message["color"] != paragraph_color)):
                if current_message["content"]:
                    messages.append({
                        "role": current_message["role"],
                        "content": self.clean_content('\n'.join(current_message["content"]))
                    })
                current_message = {"role": role, "content": [], "color": paragraph_color}
            
            # Set role and color if this is the first paragraph
            if current_message["role"] is None:
                current_message["role"] = role
                current_message["color"] = paragraph_color
            
            # Add paragraph text to current message
            current_message["content"].append(paragraph.text.strip())
        
        # Don't forget the last message
        if current_message["role"] is not None and current_message["content"]:
            messages.append({
                "role": current_message["role"],
                "content": self.clean_content('\n'.join(current_message["content"]))
            })
        
        if not messages:
            return [{
                "role": "system",
                "content": "No conversation content found in Word document. Make sure text is present and try again."
            }]
        
        return messages
    
    def _is_paragraph_highlighted(self, paragraph) -> bool:
        """
        Check if a paragraph has highlighting (indicating user message)
        """
        # Check each run in the paragraph for highlighting
        for run in paragraph.runs:
            # Check for highlight color
            if run.font.highlight_color is not None:
                # Some highlighting shows up as an enum value
                from docx.enum.text import WD_COLOR_INDEX
                if hasattr(WD_COLOR_INDEX, 'YELLOW') and run.font.highlight_color == WD_COLOR_INDEX.YELLOW:
                    return True
                if hasattr(WD_COLOR_INDEX, 'BRIGHT_GREEN') and run.font.highlight_color == WD_COLOR_INDEX.BRIGHT_GREEN:
                    return True
                # Any non-None highlight color indicates highlighting
                if str(run.font.highlight_color) != 'None':
                    return True
            
            # Also check for shading/background color
            if hasattr(run, '_element'):
                shd_elements = run._element.xpath('.//w:shd')
                if shd_elements:
                    for shd in shd_elements:
                        fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill and fill.upper() not in ['FFFFFF', 'AUTO', 'NONE', '']:
                            return True
        
        # Also check paragraph-level highlighting/shading
        if hasattr(paragraph, '_element'):
            shd_elements = paragraph._element.xpath('.//w:shd')
            if shd_elements:
                for shd in shd_elements:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill and fill.upper() not in ['FFFFFF', 'AUTO', 'NONE', '']:
                        return True
        
        return False
    
    def _get_paragraph_highlight_color(self, paragraph) -> Optional[str]:
        """
        Get the highlight color for a paragraph
        Returns the color code or None if no highlighting
        """
        # Check paragraph-level highlighting/shading first
        if hasattr(paragraph, '_element'):
            shd_elements = paragraph._element.xpath('.//w:shd')
            if shd_elements:
                for shd in shd_elements:
                    fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                    if fill and fill.upper() not in ['FFFFFF', 'AUTO', 'NONE', '']:
                        return fill.upper()
        
        # Check run-level highlighting
        for run in paragraph.runs:
            if run.font.highlight_color is not None:
                return str(run.font.highlight_color)
            
            # Check run-level shading
            if hasattr(run, '_element'):
                shd_elements = run._element.xpath('.//w:shd')
                if shd_elements:
                    for shd in shd_elements:
                        fill = shd.get('{http://schemas.openxmlformats.org/wordprocessingml/2006/main}fill')
                        if fill and fill.upper() not in ['FFFFFF', 'AUTO', 'NONE', '']:
                            return fill.upper()
        
        return None  # No highlighting found
    
    def _analyze_highlight_colors(self, doc) -> Dict:
        """
        Analyze all highlight colors in the document and their frequencies
        """
        color_counts = {}
        color_sequences = []  # Track the sequence of colors for pattern analysis
        
        for paragraph in doc.paragraphs:
            if not paragraph.text.strip():
                continue
                
            color = self._get_paragraph_highlight_color(paragraph)
            if color:
                color_counts[color] = color_counts.get(color, 0) + 1
                color_sequences.append(color)
            else:
                color_sequences.append(None)
        
        return {
            "color_counts": color_counts,
            "color_sequences": color_sequences,
            "total_colors": len(color_counts),
            "most_common_colors": sorted(color_counts.items(), key=lambda x: x[1], reverse=True)
        }
    
    def _determine_color_roles(self, color_analysis: Dict) -> Dict[str, str]:
        """
        Determine which colors represent user vs assistant messages
        
        Strategy:
        1. If we have exactly 2 colors, assume they alternate user/assistant
        2. Look at conversation patterns - shorter messages often user, longer assistant
        3. First message is typically user (common conversation pattern)
        """
        color_counts = color_analysis["color_counts"]
        color_sequences = color_analysis["color_sequences"]
        
        if len(color_counts) == 0:
            return {}  # No colors found
        elif len(color_counts) == 1:
            # Only one color - treat as user messages, None (no color) as assistant
            single_color = list(color_counts.keys())[0]
            return {single_color: "user", None: "assistant"}
        elif len(color_counts) == 2:
            # Two colors - determine which is user vs assistant
            colors = list(color_counts.keys())
            
            # Strategy: Assume first color encountered is user message
            # (conversations typically start with user input)
            first_color = None
            for color in color_sequences:
                if color is not None:
                    first_color = color
                    break
            
            if first_color:
                other_color = [c for c in colors if c != first_color][0]
                return {
                    first_color: "user",
                    other_color: "assistant",
                    None: "assistant"  # Fallback for unhighlighted text
                }
        
        # More than 2 colors - use heuristics
        # Most common color is probably assistant (longer responses)
        # Less common might be user (shorter messages)
        most_common_colors = color_analysis["most_common_colors"]
        if len(most_common_colors) >= 2:
            assistant_color = most_common_colors[0][0]  # Most frequent
            user_color = most_common_colors[1][0]       # Second most frequent
            
            return {
                user_color: "user",
                assistant_color: "assistant",
                None: "assistant"  # Fallback
            }
        
        # Fallback: treat all colors as assistant
        return {color: "assistant" for color in color_counts.keys()}
    
    def _analyze_docx_structure(self, file_path: str) -> Dict:
        """
        Analyze Word document to understand its structure and color usage
        This can help with debugging highlighting detection
        """
        try:
            doc = Document(file_path)
            
            # Basic paragraph analysis
            analysis = {
                "total_paragraphs": len(doc.paragraphs),
                "empty_paragraphs": 0,
                "paragraphs_with_color": 0,
                "paragraphs_without_color": 0
            }
            
            # Color analysis
            color_analysis = self._analyze_highlight_colors(doc)
            color_mapping = self._determine_color_roles(color_analysis)
            
            # Count paragraphs by type
            user_paragraphs = 0
            assistant_paragraphs = 0
            
            for paragraph in doc.paragraphs:
                if not paragraph.text.strip():
                    analysis["empty_paragraphs"] += 1
                    continue
                
                color = self._get_paragraph_highlight_color(paragraph)
                if color:
                    analysis["paragraphs_with_color"] += 1
                else:
                    analysis["paragraphs_without_color"] += 1
                
                # Determine role
                role = color_mapping.get(color, "assistant")
                if role == "user":
                    user_paragraphs += 1
                else:
                    assistant_paragraphs += 1
            
            # Add color and role information
            analysis.update({
                "color_analysis": color_analysis,
                "color_mapping": color_mapping,
                "user_paragraphs": user_paragraphs,
                "assistant_paragraphs": assistant_paragraphs
            })
            
            return analysis
        except Exception as e:
            return {"error": str(e)}
    
    def get_parser_name(self) -> str:
        return "Anthropic Claude Parser"
    
    def get_company_name(self) -> str:
        return "Anthropic"


class OpenAIChatParser(BaseConversationParser):
    """
    Parser for OpenAI ChatGPT exports
    Handles ChatGPT web interface exports and API conversation logs
    
    Expected formats:
    - User: [message]
      ChatGPT: [response]
    - JSON exports from ChatGPT
    - Conversation exports with code blocks, lists, etc.
    """
    
    def can_parse(self, content: str) -> bool:
        """Check for OpenAI/ChatGPT format markers"""
        markers = ["User:", "ChatGPT:", "GPT:", "GPT-4:", "GPT-3.5:"]
        json_markers = ['"role":', '"content":', '"model": "gpt']
        
        return any(marker in content for marker in markers) or \
               any(marker in content for marker in json_markers)
    
    def parse(self, content: str) -> List[Dict]:
        """
        Parse OpenAI format conversation
        TODO: Implement full parsing logic for:
        - ChatGPT web exports with formatting
        - JSON conversation exports
        - Code interpreter outputs
        - Plugin/tool usage
        - Image inputs/outputs
        - Custom GPTs with special formatting
        """
        # Placeholder implementation
        messages = []
        
        # Try JSON first
        if self._is_json_format(content):
            return self._parse_json_export(content)
        
        # Fall back to text format
        lines = content.split('\n')
        current_role = None
        current_content = []
        
        for line in lines:
            if any(line.startswith(prefix) for prefix in ["User:", "You:"]):
                if current_role:
                    messages.append({
                        "role": "user" if current_role == "User" else "assistant",
                        "content": self.clean_content('\n'.join(current_content))
                    })
                current_role = "User"
                # Find the prefix length
                prefix_len = line.index(':') + 1
                current_content = [line[prefix_len:].strip()]
            elif any(line.startswith(prefix) for prefix in ["ChatGPT:", "GPT:", "Assistant:"]):
                if current_role:
                    messages.append({
                        "role": "user" if current_role == "User" else "assistant",
                        "content": self.clean_content('\n'.join(current_content))
                    })
                current_role = "Assistant"
                prefix_len = line.index(':') + 1
                current_content = [line[prefix_len:].strip()]
            else:
                current_content.append(line)
        
        # Don't forget last message
        if current_role:
            messages.append({
                "role": "user" if current_role == "User" else "assistant",
                "content": self.clean_content('\n'.join(current_content))
            })
        
        return messages
    
    def _is_json_format(self, content: str) -> bool:
        """Check if content is JSON format"""
        content = content.strip()
        return (content.startswith('{') or content.startswith('[')) and \
               ('"role"' in content or '"messages"' in content)
    
    def _parse_json_export(self, content: str) -> List[Dict]:
        """Parse JSON format export"""
        try:
            data = json.loads(content)
            if isinstance(data, list):
                return [{"role": msg.get("role", "user"), 
                        "content": msg.get("content", "")} for msg in data]
            elif isinstance(data, dict) and 'messages' in data:
                return [{"role": msg.get("role", "user"),
                        "content": msg.get("content", "")} for msg in data['messages']]
        except:
            pass
        return []
    
    def get_parser_name(self) -> str:
        return "OpenAI ChatGPT Parser"
    
    def get_company_name(self) -> str:
        return "OpenAI"


class GeminiChatParser(BaseConversationParser):
    """
    Parser for Google Gemini (formerly Bard) chat exports
    Handles Gemini web interface exports
    
    Expected formats:
    - Various Google export formats
    - Bard legacy formats
    """
    
    def can_parse(self, content: str) -> bool:
        """Check for Gemini/Bard format markers"""
        markers = ["You:", "Gemini:", "Bard:", "Model:", "User:"]
        # Also check for Google-specific JSON structure
        google_markers = ['"prompt":', '"response":', '"conversation_id":']
        
        return any(marker in content for marker in markers) or \
               any(marker in content for marker in google_markers)
    
    def parse(self, content: str) -> List[Dict]:
        """
        Parse Gemini format conversation
        TODO: Implement full parsing logic for:
        - Gemini web exports
        - Bard legacy exports
        - Multi-modal conversations (images, code, etc.)
        - Google AI Studio exports
        - Conversation branches/variations
        """
        # Placeholder implementation
        messages = []
        
        # Basic text parsing
        lines = content.split('\n')
        current_role = None
        current_content = []
        
        for line in lines:
            if any(line.startswith(prefix) for prefix in ["You:", "User:"]):
                if current_role:
                    messages.append({
                        "role": "user" if current_role == "User" else "assistant",
                        "content": self.clean_content('\n'.join(current_content))
                    })
                current_role = "User"
                prefix_len = line.index(':') + 1
                current_content = [line[prefix_len:].strip()]
            elif any(line.startswith(prefix) for prefix in ["Gemini:", "Bard:", "Model:"]):
                if current_role:
                    messages.append({
                        "role": "user" if current_role == "User" else "assistant",
                        "content": self.clean_content('\n'.join(current_content))
                    })
                current_role = "Assistant"
                prefix_len = line.index(':') + 1
                current_content = [line[prefix_len:].strip()]
            else:
                current_content.append(line)
        
        # Don't forget last message
        if current_role:
            messages.append({
                "role": "user" if current_role == "User" else "assistant",
                "content": self.clean_content('\n'.join(current_content))
            })
        
        return messages
    
    def get_parser_name(self) -> str:
        return "Google Gemini Parser"
    
    def get_company_name(self) -> str:
        return "Google"


class GenericChatParser(BaseConversationParser):
    """
    Generic fallback parser for unknown formats
    Tries to make best guess at conversation structure
    """
    
    def can_parse(self, content: str) -> bool:
        """This parser can always try to parse as last resort"""
        return True
    
    def parse(self, content: str) -> List[Dict]:
        """
        Try to intelligently parse unknown format
        Look for conversation patterns and structure
        """
        content = self.clean_content(content)
        
        # If it looks like JSON, try to parse it
        if content.strip().startswith(('{', '[')):
            try:
                data = json.loads(content)
                if isinstance(data, list):
                    # Assume it's a message list
                    return data
                elif isinstance(data, dict):
                    # Look for common keys
                    if 'messages' in data:
                        return data['messages']
                    elif 'conversation' in data:
                        return data['conversation']
            except:
                pass
        
        # Look for common conversation patterns
        # Try to detect back-and-forth pattern
        lines = content.split('\n')
        messages = []
        
        # Simple heuristic: alternate between user and assistant
        # Look for empty lines as message boundaries
        current_content = []
        is_user = True
        
        for line in lines:
            if line.strip() == '' and current_content:
                # Empty line might indicate message boundary
                messages.append({
                    "role": "user" if is_user else "assistant",
                    "content": self.clean_content('\n'.join(current_content))
                })
                current_content = []
                is_user = not is_user
            else:
                current_content.append(line)
        
        # Add remaining content
        if current_content:
            messages.append({
                "role": "user" if is_user else "assistant",
                "content": self.clean_content('\n'.join(current_content))
            })
        
        # If we only got one message, treat entire content as system context
        if len(messages) <= 1:
            return [{"role": "system", "content": content}]
        
        return messages
    
    def get_parser_name(self) -> str:
        return "Generic Parser"
    
    def get_company_name(self) -> str:
        return "Generic"


class ConversationParserFactory:
    """
    Factory class to manage and select appropriate parsers
    """
    
    def __init__(self):
        self.parsers = [
            AnthropicChatParser(),
            OpenAIChatParser(),
            GeminiChatParser(),
            GenericChatParser()  # Always last as fallback
        ]
    
    def get_parser(self, content: str, parser_type: Optional[str] = None) -> BaseConversationParser:
        """
        Get appropriate parser for content
        
        Args:
            content: The conversation content to parse
            parser_type: Optional specific parser type to use ('anthropic', 'openai', 'gemini', 'generic')
        
        Returns:
            Appropriate parser instance
        """
        # If specific parser requested
        if parser_type:
            parser_map = {
                'anthropic': AnthropicChatParser,
                'openai': OpenAIChatParser,
                'gemini': GeminiChatParser,
                'generic': GenericChatParser
            }
            if parser_type.lower() in parser_map:
                return parser_map[parser_type.lower()]()
        
        # Auto-detect parser
        for parser in self.parsers:
            if parser.can_parse(content):
                return parser
        
        # Should never reach here as GenericParser accepts everything
        return GenericChatParser()
    
    def get_available_parsers(self) -> List[Dict]:
        """Get list of available parsers for UI"""
        return [
            {
                "id": parser.__class__.__name__.replace('ChatParser', '').lower(),
                "name": parser.get_parser_name(),
                "company": parser.get_company_name()
            }
            for parser in self.parsers
        ]